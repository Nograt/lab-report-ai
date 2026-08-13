from dataclasses import dataclass

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips




TITLE_COLUMN_WIDTHS = [
    3965,  # skład grupy / lewa część nagłówka
    1702,  # semestr
    1981,  # grupa / data
    1416,  # rok / ocena
]


TITLE_ROW_HEIGHTS = [
    1173,
    1463,
    850,
]


TITLE_FONT = "Calibri Light"
TITLE_FONT_SIZE = 12


@dataclass
class TitlePageData:
    faculty: str
    department: str
    laboratory: str

    members: list[str]

    semester: str
    group: str
    team: str
    academic_year: str

    topic: str
    execution_date: str

    grade: str = ""



def _set_run_font(
    run,
    *,
    size: float = TITLE_FONT_SIZE,
    bold: bool = False,
):
    run.font.name = TITLE_FONT
    run.font.size = Pt(size)
    run.bold = bold

    r_pr = run._element.get_or_add_rPr()

    r_fonts = r_pr.rFonts

    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)

    r_fonts.set(
        qn("w:ascii"),
        TITLE_FONT,
    )

    r_fonts.set(
        qn("w:hAnsi"),
        TITLE_FONT,
    )


def _configure_paragraph(
    paragraph,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
):
    paragraph.alignment = alignment

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def _set_paragraph_text(
    paragraph,
    text: str,
    *,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    size: float = TITLE_FONT_SIZE,
    bold: bool = False,
):
    paragraph.clear()

    _configure_paragraph(
        paragraph,
        alignment,
    )

    run = paragraph.add_run(text)

    _set_run_font(
        run,
        size=size,
        bold=bold,
    )

    return paragraph


def _add_cell_paragraph(
    cell,
    text: str,
    *,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    size: float = TITLE_FONT_SIZE,
    bold: bool = False,
):
    paragraph = cell.add_paragraph()

    _set_paragraph_text(
        paragraph,
        text,
        alignment=alignment,
        size=size,
        bold=bold,
    )

    return paragraph




def _set_cell_width(
    cell,
    width_twips: int,
):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.first_child_found_in(
        "w:tcW"
    )

    if tc_w is None:
        tc_w = OxmlElement(
            "w:tcW"
        )
        tc_pr.append(tc_w)

    tc_w.set(
        qn("w:w"),
        str(width_twips),
    )

    tc_w.set(
        qn("w:type"),
        "dxa",
    )


def _set_grid_widths(
    table,
):
    tbl_grid = table._tbl.tblGrid

    for child in list(tbl_grid):
        tbl_grid.remove(child)

    for width in TITLE_COLUMN_WIDTHS:

        grid_col = OxmlElement(
            "w:gridCol"
        )

        grid_col.set(
            qn("w:w"),
            str(width),
        )

        tbl_grid.append(
            grid_col
        )


from docx.enum.table import WD_TABLE_ALIGNMENT

def _set_cell_margins(
    cell,
    *,
    top: int = 0,
    start: int = 0,
    bottom: int = 0,
    end: int = 0,
):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_mar = tc_pr.first_child_found_in(
        "w:tcMar"
    )

    if tc_mar is None:
        tc_mar = OxmlElement(
            "w:tcMar"
        )
        tc_pr.append(tc_mar)

    for side, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():

        element = tc_mar.find(
            qn(f"w:{side}")
        )

        if element is None:
            element = OxmlElement(
                f"w:{side}"
            )
            tc_mar.append(element)

        element.set(
            qn("w:w"),
            str(value),
        )

        element.set(
            qn("w:type"),
            "dxa",
        )

def _configure_table(
    table,
):
    tbl_pr = table._tbl.tblPr


    layout = tbl_pr.first_child_found_in(
        "w:tblLayout"
    )

    if layout is None:
        layout = OxmlElement(
            "w:tblLayout"
        )

        tbl_pr.append(layout)

    layout.set(
        qn("w:type"),
        "fixed",
    )

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
 
    borders = tbl_pr.first_child_found_in(
        "w:tblBorders"
    )

    if borders is None:
        borders = OxmlElement(
            "w:tblBorders"
        )

        tbl_pr.append(borders)

    for edge in (
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ):

        element = borders.find(
            qn(f"w:{edge}")
        )

        if element is None:
            element = OxmlElement(
                f"w:{edge}"
            )

            borders.append(element)

        element.set(
            qn("w:val"),
            "single",
        )

        element.set(
            qn("w:sz"),
            "4",
        )

        element.set(
            qn("w:space"),
            "0",
        )

        element.set(
            qn("w:color"),
            "000000",
        )

 
    cell_margins = (
        tbl_pr.first_child_found_in(
            "w:tblCellMar"
        )
    )

    if cell_margins is None:
        cell_margins = OxmlElement(
            "w:tblCellMar"
        )

        tbl_pr.append(
            cell_margins
        )

    for side in (
        "left",
        "right",
    ):
        element = cell_margins.find(
            qn(f"w:{side}")
        )

        if element is None:
            element = OxmlElement(
                f"w:{side}"
            )

            cell_margins.append(
                element
            )

        element.set(
            qn("w:w"),
            "0",
        )

        element.set(
            qn("w:type"),
            "dxa",
        )



def _clear_cell(
    cell,
):
    cell.text = ""

    paragraph = cell.paragraphs[0]

    _configure_paragraph(
        paragraph
    )

    return paragraph



def add_title_page(
    document: Document,
    data: TitlePageData,
):

    table = document.add_table(
        rows=3,
        cols=4,
    )

    table.autofit = False

    _configure_table(
        table
    )

    _set_grid_widths(
        table
    )


    for row, height in zip(
        table.rows,
        TITLE_ROW_HEIGHTS,
    ):
        row.height = Twips(height)

        row.height_rule = (
            WD_ROW_HEIGHT_RULE.AT_LEAST
        )


    for row in table.rows:

        for index, width in enumerate(
            TITLE_COLUMN_WIDTHS
        ):
            _set_cell_width(
                row.cells[index],
                width,
            )



    faculty_cell = table.cell(
        0,
        0,
    )

    laboratory_cell = table.cell(
        0,
        1,
    ).merge(
        table.cell(
            0,
            3,
        )
    )

    faculty_cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    laboratory_cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


    p = _clear_cell(
        faculty_cell
    )

    _set_paragraph_text(
        p,
        data.faculty,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_cell_paragraph(
        faculty_cell,
        data.department,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


    p = _clear_cell(
        laboratory_cell
    )

    _set_paragraph_text(
        p,
        data.laboratory,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


    members_cell = table.cell(
        1,
        0,
    )
    
    _set_cell_margins(
    members_cell,
    start=120,
    end=80,
    )
    members_cell.vertical_alignment = (
    WD_CELL_VERTICAL_ALIGNMENT.TOP
    )

    semester_cell = table.cell(
        1,
        1,
    )

    group_cell = table.cell(
        1,
        2,
    )

    year_cell = table.cell(
        1,
        3,
    )



    members_cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.TOP
    )

    p = _clear_cell(
        members_cell
    )

    _set_paragraph_text(
        p,
        "Skład osobowy grupy:",
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )

    for index, member in enumerate(
        data.members,
        start=1,
    ):
        _add_cell_paragraph(
            members_cell,
            f"{index}. {member}",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )


    semester_cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    p = _clear_cell(
        semester_cell
    )

    _set_paragraph_text(
        p,
        "Semestr:",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_cell_paragraph(
        semester_cell,
        data.semester,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


    group_cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    p = _clear_cell(
        group_cell
    )

    _set_paragraph_text(
        p,
        "Grupa:",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_cell_paragraph(
        group_cell,
        data.group,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_cell_paragraph(
        group_cell,
        f"Zespół {data.team}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


    year_cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    p = _clear_cell(
        year_cell
    )

    _set_paragraph_text(
        p,
        "Rok",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_cell_paragraph(
        year_cell,
        "akademicki:",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_cell_paragraph(
        year_cell,
        data.academic_year,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


    topic_cell = table.cell(
        2,
        0,
    ).merge(
        table.cell(
            2,
            1,
        )
    )

    date_cell = table.cell(
        2,
        2,
    )

    grade_cell = table.cell(
        2,
        3,
    )
    
    grade_cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.TOP
        )

    topic_cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    p = _clear_cell(
        topic_cell
    )

    _set_paragraph_text(
        p,
        "Temat ćwiczenia:",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_cell_paragraph(
        topic_cell,
        data.topic,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


    date_cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    p = _clear_cell(
        date_cell
    )

    _set_paragraph_text(
        p,
        "Data wykonania:",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_cell_paragraph(
        date_cell,
        data.execution_date,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    p = _clear_cell(
        grade_cell
    )

    _set_paragraph_text(
        p,
        "Ocena:",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    if data.grade:
        _add_cell_paragraph(
            grade_cell,
            data.grade,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    return table
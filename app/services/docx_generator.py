from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from app.schemas.report import ReportSpecification
from app.services.excel_reader import (
    MeasurementTableData,
    get_measurement_table,
)

from app.services.word_equation import (
    append_formula_equation,
    append_substitution_equation,
    append_result_equation,
)

from app.services.report_style import (
    ReportStyle,
    DEFAULT_REPORT_STYLE,
)
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from docx.enum.section import (
    WD_ORIENT,
    WD_SECTION_START,
)

REPORT_FONT = "Calibri Light"

def calculate_image_size(
    image_path: Path,
    *,
    max_width_cm: float = 13.5,
    max_height_cm: float = 11.0,
) -> tuple[float, float]:
    """
    Zwraca rozmiar obrazka w cm.

    - zachowuje proporcje,
    - nie przekracza max_width_cm,
    - nie przekracza max_height_cm.
    """

    with Image.open(image_path) as image:
        width_px, height_px = image.size

        dpi = image.info.get("dpi", (96, 96))

        dpi_x = dpi[0] or 96
        dpi_y = dpi[1] or 96

    natural_width_cm = (
        width_px / dpi_x
    ) * 2.54

    natural_height_cm = (
        height_px / dpi_y
    ) * 2.54

    scale = min(
        1.0,
        max_width_cm / natural_width_cm,
        max_height_cm / natural_height_cm,
    )

    return (
        natural_width_cm * scale,
        natural_height_cm * scale,
    )



def set_style_font(
    style,
    font_name: str,
    font_size: Pt,
    *,
    bold: bool | None = None,
):
    style.font.name = font_name
    style.font.size = font_size

    if bold is not None:
        style.font.bold = bold

    style.font.color.rgb = RGBColor(
        0,
        0,
        0,
    )

    r_pr = style.element.get_or_add_rPr()

    r_fonts = r_pr.rFonts

    if r_fonts is not None:
        r_fonts.set(
            qn("w:ascii"),
            font_name,
        )

        r_fonts.set(
            qn("w:hAnsi"),
            font_name,
        )

        r_fonts.set(
            qn("w:eastAsia"),
            font_name,
        )


def add_page_numbers(
    document: Document,
):
    for section in document.sections:
        section.different_first_page_header_footer = False

        footer = section.footer

        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        run = paragraph.add_run()

        run.font.name = REPORT_FONT
        run.font.size = Pt(9)

        r_pr = run._element.get_or_add_rPr()

        r_fonts = r_pr.rFonts

        if r_fonts is not None:
            r_fonts.set(
                qn("w:ascii"),
                REPORT_FONT,
            )

            r_fonts.set(
                qn("w:hAnsi"),
                REPORT_FONT,
            )

   

        field_begin = OxmlElement(
            "w:fldChar"
        )

        field_begin.set(
            qn("w:fldCharType"),
            "begin",
        )

        instruction = OxmlElement(
            "w:instrText"
        )

        instruction.set(
            qn("xml:space"),
            "preserve",
        )

        instruction.text = (
            " PAGE \\* MERGEFORMAT "
        )

        field_separator = OxmlElement(
            "w:fldChar"
        )

        field_separator.set(
            qn("w:fldCharType"),
            "separate",
        )

        cached_value = OxmlElement(
            "w:t"
        )

        cached_value.text = "1"

        field_end = OxmlElement(
            "w:fldChar"
        )

        field_end.set(
            qn("w:fldCharType"),
            "end",
        )

        run._r.append(
            field_begin
        )

        run._r.append(
            instruction
        )

        run._r.append(
            field_separator
        )

        run._r.append(
            cached_value
        )

        run._r.append(
            field_end
        )

def configure_document(
    document: Document,
    style: ReportStyle,
):
   
    section = document.sections[0]

    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)



    normal = document.styles[
        "Normal"
    ]

    set_style_font(
        style=normal,
        font_name=REPORT_FONT,
        font_size=Pt(12),
    )

    normal.paragraph_format.left_indent = None
    normal.paragraph_format.right_indent = None
    normal.paragraph_format.first_line_indent = None

    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    normal.paragraph_format.line_spacing = 1.15


    heading_1 = document.styles[
        "Heading 1"
    ]

    set_style_font(
        style=heading_1,
        font_name=REPORT_FONT,
        font_size=Pt(14),
        bold=True,
    )

    heading_1.paragraph_format.left_indent = None
    heading_1.paragraph_format.right_indent = None
    heading_1.paragraph_format.first_line_indent = None

    heading_1.paragraph_format.space_before = Pt(14)
    heading_1.paragraph_format.space_after = Pt(6)

    heading_1.paragraph_format.line_spacing = 1.0

   
    heading_1.paragraph_format.keep_with_next = True

 

    heading_2 = document.styles[
        "Heading 2"
    ]

    set_style_font(
        style=heading_2,
        font_name=REPORT_FONT,
        font_size=Pt(12),
        bold=True,
    )

    heading_2.paragraph_format.left_indent = None
    heading_2.paragraph_format.right_indent = None
    heading_2.paragraph_format.first_line_indent = None

    heading_2.paragraph_format.space_before = Pt(10)
    heading_2.paragraph_format.space_after = Pt(4)

    heading_2.paragraph_format.line_spacing = 1.0
    heading_2.paragraph_format.keep_with_next = True


def set_cell_text(
    cell,
    value,
    bold: bool = False,
    font_size: float = 9,
):
    cell.text = ""

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        "" if value is None else str(value)
    )

    run.bold = bold
    run.font.name = REPORT_FONT
    run.font.size = Pt(font_size)

    run._element.get_or_add_rPr().rFonts.set(
        qn("w:ascii"),
        REPORT_FONT,
    )

    run._element.get_or_add_rPr().rFonts.set(
        qn("w:hAnsi"),
        REPORT_FONT,
    )

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


def format_number(value):
    if value is None:
        return ""

    if isinstance(value, float):
        return f"{value:.4f}".rstrip("0").rstrip(".")

    return str(value)


def add_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    borders = tbl_pr.first_child_found_in(
        "w:tblBorders"
    )

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in (
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ):
        tag = f"w:{edge}"

        element = borders.find(
            qn(tag)
        )

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(
            qn("w:val"),
            "single",
        )

        element.set(
            qn("w:sz"),
            "4",
        )

        element.set(
            qn("w:space"),
            "0",
        )

        element.set(
            qn("w:color"),
            "000000",
        )


def add_measurement_table(
    document: Document,
    table_data: MeasurementTableData,
    columns: list[str],
    table_number: int,
    title: str | None,
):
    
    widths, font_size, use_landscape = (
    choose_table_layout(
        document=document,
        table_data=table_data,
        columns=columns,
    )
)

    if use_landscape:
        start_landscape_section(document)

    caption = document.add_paragraph()

    caption.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_after = Pt(3)

    run = caption.add_run(
        f"Tabela {table_number}. "
        f"{title or table_data.title or ''}"
    )

    run.italic = True
    run.font.name = REPORT_FONT
    run.font.size = Pt(10)


    df = table_data.dataframe[
        columns
    ]

    table = document.add_table(
        rows=2,
        cols=len(columns),
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.style = "Table Grid"

    for index, column in enumerate(
        columns
    ):


        set_cell_text(
            table.rows[0].cells[index],
            column,
            bold=True,
            font_size=font_size,
        )

        unit = table_data.units.get(
            column
        )


        set_cell_text(
            table.rows[1].cells[index],
            unit or "—",
            bold=False,
            font_size=font_size,
        )


    for _, row in df.iterrows():

        cells = table.add_row().cells

        for index, column in enumerate(
            columns
        ):
            set_cell_text(
                cells[index],
                format_number(
                    row[column]
                ),
                font_size=font_size,
            )

    apply_table_column_widths(
        table=table,
        widths_cm=widths,
    )


    set_table_cell_margins(
        table
    )

    repeat_table_header(
        table.rows[0]
    )

    repeat_table_header(
        table.rows[1]
    )

    for row in table.rows:
        prevent_row_split(
            row
        )

    add_table_borders(
    table
)

    if use_landscape:
        start_portrait_section(
            document
        )
    else:
        document.add_paragraph()

def add_chart(
    document: Document,
    chart_path: Path,
    figure_number: int,
    x: str,
    y: str,
):
    width_cm, _ = calculate_image_size(
        image_path=chart_path,
        max_width_cm=13.5,
        max_height_cm=11.0,
    )


    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)

    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True

    run = paragraph.add_run()

    run.add_picture(
        str(chart_path),
        width=Cm(width_cm),
    )


    caption = document.add_paragraph()

    caption.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True

    caption_run = caption.add_run(
    f"Rys. {figure_number}. "
    f"Zależność {y} = f({x})"
)

    caption_run.italic = True
    caption_run.font.name = REPORT_FONT
    caption_run.font.size = Pt(10)


def add_example_calculations(
    document: Document,
    calculations: list[dict],
    row_index: int,
):
    if not calculations:
        return

    heading = document.add_paragraph()

    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(6)

    measurement_number = row_index + 1

    run = heading.add_run(
        f"Przykładowe obliczenia "
        f"dla pomiaru nr {measurement_number}"
    )

    run.bold = True
    run.font.name = REPORT_FONT
    run.font.size = Pt(11)

    for calculation in calculations:
        output = calculation["output"]

        calculation_heading = (
            document.add_paragraph()
        )

        calculation_heading.paragraph_format.space_before = Pt(6)
        calculation_heading.paragraph_format.space_after = Pt(3)

        calculation_heading.paragraph_format.keep_with_next = True

        run = calculation_heading.add_run(
            f"Wyznaczenie wielkości {output}:"
        )

        run.font.name = REPORT_FONT
        run.font.size = Pt(11)

        expression = calculation.get(
            "expression"
        )

        variables = calculation.get(
            "variables",
            {},
        )


        if expression is not None:


            formula_paragraph = (
                document.add_paragraph()
            )
            formula_paragraph.paragraph_format.keep_with_next = True

            formula_paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            append_formula_equation(
                paragraph=formula_paragraph,
                output=calculation["output"],
                expression=expression,
            )


            substitution_paragraph = (
                document.add_paragraph()
            )
            substitution_paragraph.paragraph_format.keep_with_next = True

            substitution_paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            append_substitution_equation(
                paragraph=substitution_paragraph,
                output=calculation["output"],
                expression=expression,
                variables=variables,
            )

            result_paragraph = (
                document.add_paragraph()
            )
            result_paragraph.paragraph_format.keep_with_next = False

            result_paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            append_result_equation(
                paragraph=result_paragraph,
                output=calculation["output"],
                result=calculation["result"],
                unit=calculation.get("unit"),
            )

            continue


        for key in (
            "formula_latex",
            "substitution_latex",
            "result_latex",
        ):

            value = calculation.get(key)

            if not value:
                continue

            value = (
                value
                .replace(
                    "\\cdot",
                    "·",
                )
                .replace(
                    "\\times",
                    "×",
                )
            )

            paragraph = (
                document.add_paragraph()
            )

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            run = paragraph.add_run(
                value
            )

            run.font.name = (
                REPORT_FONT
            )

            run.font.size = Pt(11)



from app.services.title_page import (
    TitlePageData,
    add_title_page,
)

def estimate_text_width_cm(
    value,
    font_size: float = 9,
) -> float:
   

    text = (
        ""
        if value is None
        else str(value)
    )

    width_units = 0.0

    narrow_chars = set(
        "il1.,:;|' "
    )

    wide_chars = set(
        "MW@%&"
    )

    for char in text:

        if char in narrow_chars:
            width_units += 0.30

        elif char in wide_chars:
            width_units += 0.90

        else:
            width_units += 0.55

    width_cm = (
        width_units
        * font_size
        * 0.0353
    )

    return width_cm

def calculate_natural_column_widths(
    table_data: MeasurementTableData,
    columns: list[str],
    font_size: float,
) -> list[float]:

    df = table_data.dataframe[
        columns
    ]

    widths = []

    horizontal_padding_cm = 0.4

    for column in columns:

        values = [
            column,
            table_data.units.get(
                column
            ) or "—",
        ]

        values.extend(
            format_number(value)
            for value
            in df[column].tolist()
        )

        widest_content = max(
            estimate_text_width_cm(
                value,
                font_size=font_size,
            )
            for value in values
        )

        width = (
            widest_content
            + horizontal_padding_cm
        )

        width = max(
            width,
            1.1,
        )

        width = min(
            width,
            4.5,
        )

        widths.append(
            width
        )

    return widths

def scale_widths_to_limit(
    widths: list[float],
    max_width_cm: float,
) -> list[float]:

    total_width = sum(
        widths
    )

    if total_width <= max_width_cm:
        return widths

    scale = (
        max_width_cm
        / total_width
    )

    return [
        width * scale
        for width in widths
    ]
    
def choose_table_layout(
    document: Document,
    table_data: MeasurementTableData,
    columns: list[str],
) -> tuple[list[float], float, bool]:

    section = document.sections[-1]


    portrait_width_cm = (
        section.page_width.cm
        - section.left_margin.cm
        - section.right_margin.cm
        - 0.4
    )

    landscape_width_cm = (
        section.page_height.cm
        - section.left_margin.cm
        - section.right_margin.cm
        - 0.4
    )



    widths_9 = (
        calculate_natural_column_widths(
            table_data=table_data,
            columns=columns,
            font_size=9,
        )
    )

    total_9 = sum(
        widths_9
    )

    if total_9 <= portrait_width_cm:
        return (
            widths_9,
            9,
            False,
        )

    widths_8 = (
        calculate_natural_column_widths(
            table_data=table_data,
            columns=columns,
            font_size=8,
        )
    )

    total_8 = sum(
        widths_8
    )

    compression = (
        portrait_width_cm
        / total_8
    )

    if compression >= 0.85:

        return (
            scale_widths_to_limit(
                widths=widths_8,
                max_width_cm=portrait_width_cm,
            ),
            8,
            False,
        )


    if total_9 <= landscape_width_cm:
        return (
            widths_9,
            9,
            True,
        )

    return (
        scale_widths_to_limit(
            widths=widths_8,
            max_width_cm=landscape_width_cm,
        ),
        8,
        True,
    )
    
def start_landscape_section(
    document: Document,
):
    previous_section = (
        document.sections[-1]
    )

    previous_width = (
        previous_section.page_width
    )

    previous_height = (
        previous_section.page_height
    )

    new_section = document.add_section(
        WD_SECTION_START.NEW_PAGE
    )

    new_section.orientation = (
        WD_ORIENT.LANDSCAPE
    )

    new_section.page_width = (
        previous_height
    )

    new_section.page_height = (
        previous_width
    )

    new_section.top_margin = (
        previous_section.top_margin
    )

    new_section.bottom_margin = (
        previous_section.bottom_margin
    )

    new_section.left_margin = (
        previous_section.left_margin
    )

    new_section.right_margin = (
        previous_section.right_margin
    )

    return new_section


def start_portrait_section(
    document: Document,
):
    previous_section = (
        document.sections[-1]
    )

    previous_width = (
        previous_section.page_width
    )

    previous_height = (
        previous_section.page_height
    )

    new_section = document.add_section(
        WD_SECTION_START.NEW_PAGE
    )

    new_section.orientation = (
        WD_ORIENT.PORTRAIT
    )

    new_section.page_width = (
        previous_height
    )

    new_section.page_height = (
        previous_width
    )

    new_section.top_margin = (
        previous_section.top_margin
    )

    new_section.bottom_margin = (
        previous_section.bottom_margin
    )

    new_section.left_margin = (
        previous_section.left_margin
    )

    new_section.right_margin = (
        previous_section.right_margin
    )

    return new_section

def apply_table_column_widths(
    table,
    widths_cm: list[float],
):
    table.autofit = False

    tbl_pr = table._tbl.tblPr


    total_width = sum(
        widths_cm
    )

    tbl_w = tbl_pr.first_child_found_in(
        "w:tblW"
    )

    if tbl_w is None:
        tbl_w = OxmlElement(
            "w:tblW"
        )

        tbl_pr.append(
            tbl_w
        )

    tbl_w.set(
        qn("w:type"),
        "dxa",
    )

    tbl_w.set(
        qn("w:w"),
        str(
            int(
                Cm(total_width).twips
            )
        ),
    )


    grid_columns = list(
        table._tbl.tblGrid
    )

    for index, width_cm in enumerate(
        widths_cm
    ):
        width = Cm(
            width_cm
        )

        if index < len(grid_columns):

            grid_columns[index].set(
                qn("w:w"),
                str(
                    int(width.twips)
                ),
            )

        for row in table.rows:

            cell = row.cells[
                index
            ]

            cell.width = width

            tc_pr = (
                cell._tc.get_or_add_tcPr()
            )

            tc_w = (
                tc_pr.first_child_found_in(
                    "w:tcW"
                )
            )

            if tc_w is None:
                tc_w = OxmlElement(
                    "w:tcW"
                )

                tc_pr.append(
                    tc_w
                )

            tc_w.set(
                qn("w:type"),
                "dxa",
            )

            tc_w.set(
                qn("w:w"),
                str(
                    int(width.twips)
                ),
            )
            
def set_table_cell_margins(
    table,
    left: int = 110,
    right: int = 110,
    top: int = 40,
    bottom: int = 40,
):
    tbl_pr = table._tbl.tblPr

    margins = (
        tbl_pr.first_child_found_in(
            "w:tblCellMar"
        )
    )

    if margins is None:
        margins = OxmlElement(
            "w:tblCellMar"
        )

        tbl_pr.append(
            margins
        )

    values = {
        "top": top,
        "left": left,
        "bottom": bottom,
        "right": right,
    }

    for side, value in values.items():

        element = margins.find(
            qn(
                f"w:{side}"
            )
        )

        if element is None:
            element = OxmlElement(
                f"w:{side}"
            )

            margins.append(
                element
            )

        element.set(
            qn("w:w"),
            str(value),
        )

        element.set(
            qn("w:type"),
            "dxa",
        )
        
def repeat_table_header(
    row,
):
    tr_pr = row._tr.get_or_add_trPr()

    tbl_header = OxmlElement(
        "w:tblHeader"
    )

    tbl_header.set(
        qn("w:val"),
        "true",
    )

    tr_pr.append(
        tbl_header
    )
    
def prevent_row_split(
    row,
):
    tr_pr = row._tr.get_or_add_trPr()

    cant_split = OxmlElement(
        "w:cantSplit"
    )

    tr_pr.append(
        cant_split
    )
            
def generate_report_docx(
    report_dir: Path,
    state: dict,
    tables: list[MeasurementTableData],
    style: ReportStyle = DEFAULT_REPORT_STYLE,
) -> Path:

    specification = (
        ReportSpecification.model_validate(
            state["specification"]
        )
    )

    report_text = state[
        "report_text"
    ]

    document = Document()

    configure_document(
        document=document,
        style=style,
    )
    
    add_page_numbers(
    document
)


    title_page_data = TitlePageData(
        faculty=(
            "WYDZIAŁ ELEKTROTECHNIKI I "
            "INFORMATYKI PL"
        ),
        department=(
            "Katedra Napędów i Maszyn "
            "Elektrycznych"
        ),
        laboratory=(
            "Laboratorium elektromaszynowych "
            "układów wykonawczych"
        ),
        members=[
            "Wojciech Targoński",
            "Jan Walczyński",
            "Marcin Piela",
            "Miłosz Pietrak",
        ],
        semester="4",
        group="IZI 4.2/3",
        team="1",
        academic_year="2025/2026",
        topic=specification.report_title,
        execution_date="26.05.2026",
        grade="",
    )

    add_title_page(
        document=document,
        data=title_page_data,
    )


    spacer = document.add_paragraph()

    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(10)

    document.add_heading(
        "1. Cel ćwiczenia",
        level=1,
    )

    document.add_paragraph(
        report_text["purpose"]
    )

   

    theory = report_text.get(
    "theory"
)

    if theory:
        document.add_heading(
        "2. Podstawy teoretyczne",
        level=1,
    )

        document.add_paragraph(
        theory
    )

        results_number = 3
        conclusions_number = 4

    else:
        results_number = 2
        conclusions_number = 3

    document.add_heading(
        f"{results_number}. Opracowanie wyników",
        level=1,
    )

    text_by_section_id = {
        section["section_id"]: section
        for section in report_text["sections"]
    }

    examples_by_section_id = {
        item["section_id"]: item
        for item in state.get(
            "example_calculations",
            [],
        )
    }

    charts_by_figure_id = {
        chart["figure_id"]: chart
        for chart in state.get(
            "charts",
            [],
        )
    }
    
    setup_images_by_id = {
    image["image_id"]: image
    for image in state.get(
        "setup_images",
        [],
    )
}

    section_setup_images = state.get(
    "section_setup_images",
    {},
)

    shown_setup_image_ids = set()

    figure_number = 1
    table_number = 1


    for section_index, section in enumerate(
        specification.sections,
        start=1,
    ):

        document.add_heading(
            (
                f"{results_number}.{section_index}. "
                f"{section.title}"
            ),
            level=2,
        )

        section_text = (
            text_by_section_id.get(
                section.section_id,
                {},
            )
        )


        if (
            section.include_description
            and section_text.get("description")
        ):
            document.add_paragraph(
                section_text["description"]
            )
            

        setup_image_id = section_setup_images.get(
            str(section.section_id)
        )

        if (
            setup_image_id
            and setup_image_id
            not in shown_setup_image_ids
        ):
            setup_image = setup_images_by_id.get(
                setup_image_id
            )

            if setup_image is None:
                raise ValueError(
                    f"Unknown setup image "
                    f"'{setup_image_id}' "
                    f"for section "
                    f"{section.section_id}."
                )

            setup_image_path = (
                report_dir
                / "setup"
                / setup_image["filename"]
            )

            add_setup_image(
                document=document,
                image_path=setup_image_path,
                figure_number=figure_number,
                caption=setup_image.get(
                    "caption"
                ),
            )

            figure_number += 1

            shown_setup_image_ids.add(
                setup_image_id
            )

       
       

        if section.table is not None:

            table_data = get_measurement_table(
                tables=tables,
                table_id=section.table_id,
            )

            add_measurement_table(
                document=document,
                table_data=table_data,
                columns=section.table.columns,
                table_number=table_number,
                title=section.table.title,
            )

            table_number += 1


        section_examples = (
            examples_by_section_id.get(
                section.section_id
            )
        )

        if section_examples:

            add_example_calculations(
                document=document,
                calculations=section_examples[
                    "calculations"
                ],
                row_index=section_examples[
                    "row_index"
                ],
            )

        for figure_id in section.chart_figure_ids:

            chart = charts_by_figure_id.get(
                figure_id
            )

            if chart is None:
                continue

            chart_path = (
                report_dir
                / "charts"
                / f"figure_{figure_id}.png"
            )

            if not chart_path.exists():
                continue

            add_chart(
                document=document,
                chart_path=chart_path,
                figure_number=figure_number,
                x=chart["x"],
                y=chart["y"],
            )

            
            figure_number += 1


        if (
            section.include_analysis
            and section_text.get("analysis")
        ):
            analysis_heading = (
                document.add_paragraph()
            )
            
            analysis_heading.paragraph_format.keep_with_next = True

            analysis_heading.paragraph_format.space_before = Pt(6)
            analysis_heading.paragraph_format.space_after = Pt(3)

            run = analysis_heading.add_run(
                "Analiza wyników"
            )

            run.bold = True
            run.font.name = style.body_font
            run.font.size = Pt(11)

            document.add_paragraph(
                section_text["analysis"]
            )

    document.add_heading(
        f"{conclusions_number}. Wnioski",
        level=1,
    )

    document.add_paragraph(
        report_text["conclusions"]
    )


    output_path = (
        report_dir
        / "report.docx"
    )

    document.save(
        output_path
    )

    return output_path

def add_setup_image(
    document: Document,
    image_path: Path,
    figure_number: int,
    caption: str | None = None,
):
    if not image_path.exists():
        raise ValueError(
            f"Setup image does not exist: "
            f"{image_path}"
        )

    width_cm, _ = calculate_image_size(
        image_path=image_path,
        max_width_cm=13.5,
        max_height_cm=11.0,
    )


    image_paragraph = document.add_paragraph()

    image_paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    image_paragraph.paragraph_format.space_before = Pt(6)
    image_paragraph.paragraph_format.space_after = Pt(3)

    image_paragraph.paragraph_format.keep_together = True

    if caption:
        image_paragraph.paragraph_format.keep_with_next = True

    run = image_paragraph.add_run()

    run.add_picture(
        str(image_path),
        width=Cm(width_cm),
    )


    if caption:
        caption_paragraph = document.add_paragraph()

        caption_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        caption_paragraph.paragraph_format.space_before = Pt(0)
        caption_paragraph.paragraph_format.space_after = Pt(8)
        caption_paragraph.paragraph_format.keep_together = True

        run = caption_paragraph.add_run(
            f"Rys. {figure_number}. {caption}"
        )

        run.italic = True
        run.font.name = REPORT_FONT
        run.font.size = Pt(10)